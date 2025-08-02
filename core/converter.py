"""
File converter for Thomson PDF Generator
Converts various file formats to PDF without external APIs
"""
import os
import io
from pathlib import Path
from typing import Optional, Dict, Any, Callable
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from PIL import Image as PILImage
import docx
import openpyxl
from utils.file_utils import FileUtils


class FileConverter:
    """Handles conversion of various file formats to PDF"""
    
    def __init__(self, progress_callback: Optional[Callable[[str, int], None]] = None):
        """
        Initialize converter with optional progress callback
        
        Args:
            progress_callback: Function to call with (status_message, progress_percentage)
        """
        self.progress_callback = progress_callback
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Setup custom paragraph styles"""
        # Title style
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#1f538d')
        ))
        
        # Subtitle style
        self.styles.add(ParagraphStyle(
            name='CustomSubtitle',
            parent=self.styles['Heading2'],
            fontSize=14,
            spaceAfter=20,
            textColor=colors.HexColor('#14a085')
        ))
    
    def _update_progress(self, message: str, percentage: int):
        """Update progress if callback is provided"""
        if self.progress_callback:
            self.progress_callback(message, percentage)
    
    def convert_to_pdf(self, input_path: str, output_path: str) -> bool:
        """
        Convert file to PDF based on its type
        
        Args:
            input_path: Path to input file
            output_path: Path for output PDF
            
        Returns:
            True if conversion successful, False otherwise
        """
        try:
            if not FileUtils.validate_file_exists(input_path):
                raise FileNotFoundError(f"Input file not found: {input_path}")
            
            file_type = FileUtils.get_file_type(input_path)
            if not file_type:
                raise ValueError(f"Unsupported file type: {input_path}")
            
            self._update_progress("Starting conversion...", 10)
            
            # Route to appropriate converter based on file type
            if file_type == 'text/plain':
                return self._convert_text_to_pdf(input_path, output_path)
            elif file_type == 'application/msword':
                return self._convert_doc_to_pdf(input_path, output_path)
            elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return self._convert_docx_to_pdf(input_path, output_path)
            elif file_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                return self._convert_xlsx_to_pdf(input_path, output_path)
            elif file_type.startswith('image/'):
                return self._convert_image_to_pdf(input_path, output_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            self._update_progress(f"Conversion failed: {str(e)}", 0)
            return False
    
    def _convert_text_to_pdf(self, input_path: str, output_path: str) -> bool:
        """Convert text file to PDF"""
        try:
            self._update_progress("Reading text file...", 30)
            
            # Detect encoding and read file
            encoding = FileUtils.detect_text_encoding(input_path)
            with open(input_path, 'r', encoding=encoding) as file:
                content = file.read()
            
            self._update_progress("Creating PDF...", 60)
            
            # Create PDF document
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            story = []
            
            # Add title
            filename = Path(input_path).stem
            title = Paragraph(f"Document: {filename}", self.styles['CustomTitle'])
            story.append(title)
            story.append(Spacer(1, 20))
            
            # Process content line by line
            lines = content.split('\n')
            for line in lines:
                if line.strip():
                    para = Paragraph(line, self.styles['Normal'])
                    story.append(para)
                else:
                    story.append(Spacer(1, 6))
            
            doc.build(story)
            self._update_progress("Text conversion completed!", 100)
            return True
            
        except Exception as e:
            self._update_progress(f"Text conversion failed: {str(e)}", 0)
            return False
    
    def _convert_docx_to_pdf(self, input_path: str, output_path: str) -> bool:
        """Convert DOCX file to PDF"""
        try:
            self._update_progress("Reading DOCX file...", 30)
            
            # Read DOCX document
            doc_docx = docx.Document(input_path)
            
            self._update_progress("Converting to PDF...", 60)
            
            # Create PDF document
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            story = []
            
            # Add title
            filename = Path(input_path).stem
            title = Paragraph(f"Document: {filename}", self.styles['CustomTitle'])
            story.append(title)
            story.append(Spacer(1, 20))
            
            # Process paragraphs
            for paragraph in doc_docx.paragraphs:
                if paragraph.text.strip():
                    # Determine style based on paragraph style
                    style_name = 'Normal'
                    if paragraph.style.name.startswith('Heading'):
                        style_name = 'CustomSubtitle'
                    
                    para = Paragraph(paragraph.text, self.styles[style_name])
                    story.append(para)
                    story.append(Spacer(1, 6))
            
            # Process tables
            for table in doc_docx.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data.append(row_data)
                
                if table_data:
                    pdf_table = Table(table_data)
                    pdf_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 14),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(pdf_table)
                    story.append(Spacer(1, 12))
            
            doc.build(story)
            self._update_progress("DOCX conversion completed!", 100)
            return True
            
        except Exception as e:
            self._update_progress(f"DOCX conversion failed: {str(e)}", 0)
            return False
    
    def _convert_doc_to_pdf(self, input_path: str, output_path: str) -> bool:
        """Convert DOC file to PDF (basic conversion)"""
        try:
            self._update_progress("DOC file detected...", 30)
            
            # For .doc files, we'll create a simple PDF with a message
            # indicating that full .doc support requires additional libraries
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            story = []
            
            filename = Path(input_path).stem
            title = Paragraph(f"Document: {filename}", self.styles['CustomTitle'])
            story.append(title)
            story.append(Spacer(1, 20))
            
            message = Paragraph(
                "This is a .doc file. For full .doc file support, please save your document as .docx format. "
                "Basic file information has been preserved in this PDF.",
                self.styles['Normal']
            )
            story.append(message)
            story.append(Spacer(1, 20))
            
            # Add file information
            file_size = FileUtils.get_file_size(input_path)
            info_text = f"Original file: {Path(input_path).name}<br/>File size: {file_size} bytes"
            info = Paragraph(info_text, self.styles['Normal'])
            story.append(info)
            
            doc.build(story)
            self._update_progress("DOC conversion completed!", 100)
            return True
            
        except Exception as e:
            self._update_progress(f"DOC conversion failed: {str(e)}", 0)
            return False
    
    def _convert_xlsx_to_pdf(self, input_path: str, output_path: str) -> bool:
        """Convert XLSX file to PDF"""
        try:
            self._update_progress("Reading Excel file...", 30)
            
            # Read Excel workbook
            workbook = openpyxl.load_workbook(input_path, data_only=True)
            
            self._update_progress("Converting sheets to PDF...", 60)
            
            # Create PDF document
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            story = []
            
            # Add title
            filename = Path(input_path).stem
            title = Paragraph(f"Workbook: {filename}", self.styles['CustomTitle'])
            story.append(title)
            story.append(Spacer(1, 20))
            
            # Process each worksheet
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]
                
                # Add sheet title
                sheet_title = Paragraph(f"Sheet: {sheet_name}", self.styles['CustomSubtitle'])
                story.append(sheet_title)
                story.append(Spacer(1, 12))
                
                # Get data from sheet
                data = []
                for row in worksheet.iter_rows(values_only=True):
                    # Convert None values to empty strings and limit row length
                    row_data = [str(cell) if cell is not None else '' for cell in row[:10]]  # Limit to 10 columns
                    if any(cell.strip() for cell in row_data):  # Only add non-empty rows
                        data.append(row_data)
                    
                    if len(data) >= 50:  # Limit to 50 rows per sheet
                        break
                
                if data:
                    # Create table
                    pdf_table = Table(data)
                    pdf_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 8),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(pdf_table)
                else:
                    empty_msg = Paragraph("(Empty sheet)", self.styles['Normal'])
                    story.append(empty_msg)
                
                story.append(Spacer(1, 20))
            
            doc.build(story)
            self._update_progress("Excel conversion completed!", 100)
            return True
            
        except Exception as e:
            self._update_progress(f"Excel conversion failed: {str(e)}", 0)
            return False
    
    def _convert_image_to_pdf(self, input_path: str, output_path: str) -> bool:
        """Convert image file to PDF"""
        try:
            self._update_progress("Processing image...", 30)
            
            # Open and process image
            with PILImage.open(input_path) as img:
                # Convert to RGB if necessary
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                self._update_progress("Creating PDF...", 60)
                
                # Create PDF document
                doc = SimpleDocTemplate(output_path, pagesize=letter)
                story = []
                
                # Add title
                filename = Path(input_path).stem
                title = Paragraph(f"Image: {filename}", self.styles['CustomTitle'])
                story.append(title)
                story.append(Spacer(1, 20))
                
                # Calculate image size to fit page
                page_width, page_height = letter
                margin = 72  # 1 inch margin
                available_width = page_width - 2 * margin
                available_height = page_height - 200  # Leave space for title and margins
                
                img_width, img_height = img.size
                aspect_ratio = img_width / img_height
                
                if img_width > available_width or img_height > available_height:
                    if aspect_ratio > 1:
                        # Landscape image
                        new_width = available_width
                        new_height = new_width / aspect_ratio
                    else:
                        # Portrait image
                        new_height = available_height
                        new_width = new_height * aspect_ratio
                else:
                    new_width, new_height = img_width, img_height
                
                # Save image to temporary buffer
                img_buffer = io.BytesIO()
                img.save(img_buffer, format='PNG')
                img_buffer.seek(0)
                
                # Add image to PDF
                pdf_image = Image(img_buffer, width=new_width, height=new_height)
                story.append(pdf_image)
                
                doc.build(story)
                self._update_progress("Image conversion completed!", 100)
                return True
                
        except Exception as e:
            self._update_progress(f"Image conversion failed: {str(e)}", 0)
            return False